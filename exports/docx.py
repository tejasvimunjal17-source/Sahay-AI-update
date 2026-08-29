"""
exports/docx.py
------------------
PHASE 6 IMPLEMENTATION.

Renders a exports._shared.ReportData into a DOCX using python-docx. Same
content, structure, and terminology as exports/pdf.py — both read from
the identical ReportData shape, so nothing here re-decides what a
Wellness Reflection Report contains (see exports/_shared.py for that).
python-docx is imported lazily inside the function, matching this
codebase's established pattern, even though (unlike fpdf2) it happens to
already be installed in this environment.
"""

from __future__ import annotations

import io

from exports._shared import ReportData


class DocxExportError(RuntimeError):
    """User-safe error message — never contains raw library internals."""


def build_docx_report(data: ReportData) -> bytes:
    """Returns DOCX bytes for the given (already-bounded, already-shaped)
    report data. Raises DocxExportError with a friendly message on any
    failure — never raises a raw python-docx exception to the caller."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise DocxExportError(
            "DOCX export isn't available right now — the required library isn't installed."
        ) from exc

    try:
        doc = Document()
        _header(doc, data, Pt, RGBColor)
        _summary_section(doc, data, Pt)

        if not data.has_any_data:
            p = doc.add_paragraph("No wellness activity was recorded in this period.")
            p.italic = True
        else:
            _mood_section(doc, data, Pt)
            _conversations_section(doc, data, Pt, RGBColor)
            _distribution_section(doc, data)

        _disclaimer_section(doc, data, Pt, RGBColor)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except DocxExportError:
        raise
    except Exception as exc:  # noqa: BLE001 - any python-docx-internal failure becomes a friendly error
        raise DocxExportError("Couldn't generate the DOCX report right now. Please try again.") from exc


def _header(doc, data: ReportData, Pt, RGBColor) -> None:
    title = doc.add_heading("Sahay AI", level=1)
    doc.add_heading("Wellness Reflection Report", level=2)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Generated {data.generated_at}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    if data.display_name:
        p = doc.add_paragraph()
        r = p.add_run(f"For: {data.display_name}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    period = doc.add_paragraph()
    period_run = period.add_run(f"Period: {data.period_start} - {data.period_end} ({data.period_days} days)")
    period_run.font.size = Pt(10)
    period_run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)


def _summary_section(doc, data: ReportData, Pt) -> None:
    doc.add_heading("Summary", level=3)
    lines = [
        f"Conversations in this period: {len(data.conversations_summary)}",
        f"Mood entries recorded: {len(data.mood_events)}",
        f"Relaxation activities completed: {data.activities_completed}",
    ]
    if data.stress_avg is not None:
        lines.append(f"Average recorded stress: {data.stress_avg}/5")
    if data.energy_avg is not None:
        lines.append(f"Average recorded energy: {data.energy_avg}/5")
    if data.sleep_avg is not None:
        lines.append(f"Average recorded sleep quality: {data.sleep_avg}/5")
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def _mood_section(doc, data: ReportData, Pt) -> None:
    if not data.mood_events:
        return
    doc.add_heading("Mood History", level=3)
    for m in data.mood_events:
        scales = []
        if m.get("stress") is not None:
            scales.append(f"Stress {m['stress']}/5")
        if m.get("energy") is not None:
            scales.append(f"Energy {m['energy']}/5")
        if m.get("sleep") is not None:
            scales.append(f"Sleep {m['sleep']}/5")
        scale_text = (" - " + ", ".join(scales)) if scales else ""
        line = f"{m['date']}  {m.get('mood', 'Neutral')} ({m.get('source', '')}){scale_text}"
        doc.add_paragraph(line, style="List Bullet")
        if m.get("note"):
            note_p = doc.add_paragraph(f"   Note: {m['note']}")
            note_p.runs[0].italic = True


def _conversations_section(doc, data: ReportData, Pt, RGBColor) -> None:
    if not data.conversations_summary:
        return
    doc.add_heading("Conversations", level=3)
    note_p = doc.add_paragraph("Titles and message counts only — full conversation content is not included in this report.")
    note_p.runs[0].font.size = Pt(9)
    note_p.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    for c in data.conversations_summary:
        doc.add_paragraph(f"{c['date']}  {c['title']} ({c['message_count']} messages)", style="List Bullet")


def _distribution_section(doc, data: ReportData) -> None:
    if not data.mood_distribution:
        return
    doc.add_heading("Approximate Mood Distribution", level=3)
    for mood, count in sorted(data.mood_distribution.items(), key=lambda kv: -kv[1]):
        doc.add_paragraph(f"{mood}: {count}", style="List Bullet")


def _disclaimer_section(doc, data: ReportData, Pt, RGBColor) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(data.disclaimer)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
